"""
Backend logic for the Legal Document Simplifier AI.

- Uses OpenRouter via langchain_openai.ChatOpenAI
- Uses FAISS for in-memory vector storage (no sqlite)
- Includes date utilities and safe wrappers for robustness
"""

import json
import os
import difflib
import re
import io
from typing import Tuple, Dict, List, Any
from datetime import datetime, timedelta
from dateutil.parser import parse as parse_date

from dotenv import load_dotenv
load_dotenv()

# -------------------------------------------------------------------------
# LangChain / Embeddings / Vector DB / LLM
# -------------------------------------------------------------------------
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain.chains.base import LLMChain
from langchain_openai import ChatOpenAI


# File parsing
import pdfplumber
import docx

# -------------------------------------------------------------------------
# Config / Env
# -------------------------------------------------------------------------
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
MODEL_ID = os.getenv("MODEL_ID", "nvidia/nemotron-nano-9b-v2:free")

# Helpful debug print at import time (will show up in logs)
print("=== BACKEND STARTUP ===")
print(f"Model ID from env: {MODEL_ID}")
print(f"OPENROUTER_API_KEY present: {'YES' if OPENROUTER_API_KEY else 'NO'}")
print("=======================")

# -------------------------------------------------------------------------
# LLM client factory (OpenRouter)
# -------------------------------------------------------------------------
def _get_llm(model_name: str = None, temperature: float = 0.0):
    """Return a ChatOpenAI client configured for OpenRouter and log key/base status."""
    if model_name is None:
        model_name = MODEL_ID

    # Debug logs to help diagnose 401/auth issues on deploy
    try:
        key_present = bool(OPENROUTER_API_KEY and len(OPENROUTER_API_KEY) > 8)
    except Exception:
        key_present = False

    print(f"🔧 _get_llm() -> model: {model_name}")
    print(f"🔧 OPENROUTER key loaded? {'YES' if key_present else 'NO'}")
    print("🔧 openai_api_base: https://openrouter.ai/api/v1 (requests will be routed there)")

    return ChatOpenAI(
        model=model_name,
        openai_api_key=OPENROUTER_API_KEY,
        openai_api_base="https://openrouter.ai/api/v1",
        temperature=temperature,
    )

# -------------------------------------------------------------------------
# File parsing
# -------------------------------------------------------------------------
def parse_file(uploaded_file) -> str:
    """Parse uploaded pdf/docx file-like object and return extracted text."""
    fname = uploaded_file.name.lower()
    content = uploaded_file.read()
    text = ""
    if fname.endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = "\n\n".join([page.extract_text() or "" for page in pdf.pages])
    elif fname.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        text = "\n\n".join([p.text for p in doc.paragraphs])
    else:
        # fallback for plain text uploads
        text = content.decode("utf-8", errors="ignore")
    return text

# -------------------------------------------------------------------------
# Embeddings & Vector store (FAISS, in-memory)
# -------------------------------------------------------------------------
def _get_embedding_model(embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
    """Return a HuggingFaceEmbeddings instance."""
    return HuggingFaceEmbeddings(model_name=embedding_model_name)

# Global in-memory stores (per-process)
VECTOR_STORES: Dict[str, FAISS] = {}

def chunk_and_store(
    text: str,
    collection_name: str = "legal_docs",
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
):
    """
    Split text into chunks, embed, and store in an in-memory FAISS vector store.
    Returns the vector store object.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_text(text)
    docs = [
        Document(page_content=c, metadata={"chunk_id": str(i), "source": collection_name})
        for i, c in enumerate(chunks)
    ]

    embedder = _get_embedding_model(embedding_model_name)
    vectordb = FAISS.from_documents(docs, embedder)
    VECTOR_STORES[collection_name] = vectordb
    print(f"📚 Stored {len(chunks)} chunks into FAISS collection '{collection_name}'")
    return vectordb

# -------------------------------------------------------------------------
# Core RAG functions
# -------------------------------------------------------------------------
def semantic_search(
    query: str,
    collection_name: str = "legal_docs",
    top_k: int = 5,
    embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
    model_name: str = None,
) -> str:
    """Retrieve relevant chunks and call LLM to answer in plain English."""
    if collection_name not in VECTOR_STORES:
        return "⚠️ No vector store found. Please upload and process a document first."

    vectordb = VECTOR_STORES[collection_name]
    docs = vectordb.similarity_search(query, k=top_k)

    # Build context from retrieved chunks
    ctxs = [f"--- chunk {i} ---\n{d.page_content.strip()[:2000]}" for i, d in enumerate(docs)]
    context_text = "\n\n".join(ctxs)

    llm = _get_llm(model_name=model_name, temperature=0.0)
    prompt_template = """You are a legal assistant. Use these retrieved chunks to answer:

RETRIEVED CHUNKS:
{context}

QUESTION:
{question}

Answer clearly in plain English for a non-lawyer.
"""
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = LLMChain(llm=llm, prompt=prompt)
    return chain.run({"context": context_text, "question": query})

def summarize_text(text: str, model_name: str = None) -> str:
    """Summarize a document in plain language (LLM call)."""
    llm = _get_llm(model_name=model_name, temperature=0.0)
    prompt_template = """Summarize the document in plain language.
Highlight purpose, key obligations, important dates, and immediate actions.

DOCUMENT:
{doc}"""
    prompt = PromptTemplate(template=prompt_template, input_variables=["doc"])
    chain = LLMChain(llm=llm, prompt=prompt)
    return chain.run({"doc": text})

def analyze_document_for_risks(text: str, model_name: str = None) -> Tuple[Dict[str, List[str]], List[str]]:
    """Ask LLM to identify risks and obligations and return parsed JSON (fallback-safe)."""
    llm = _get_llm(model_name=model_name, temperature=0.0)
    prompt_template = """Analyze this document and identify risks and obligations.

DOCUMENT:
{doc}

Respond in JSON:
{
  "High": ["..."],
  "Medium": ["..."],
  "Low": ["..."],
  "Obligations": ["..."]
}"""
    prompt = PromptTemplate(template=prompt_template, input_variables=["doc"])
    chain = LLMChain(llm=llm, prompt=prompt)
    raw = chain.run({"doc": text[:8000]})
    try:
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = "\n".join(cleaned.splitlines()[1:-1])
        parsed = json.loads(re.search(r"\{.*\}", cleaned, re.DOTALL).group(0))
    except Exception as e:
        print("⚠️ Risk parsing failed:", e)
        parsed = {"High": [], "Medium": [], "Low": [], "Obligations": []}
    return {k: parsed.get(k, []) for k in ["High", "Medium", "Low"]}, parsed.get("Obligations", [])

def safe_analyze_document_for_risks(text: str, model_name: str = None) -> Tuple[Dict[str, List[str]], List[str]]:
    """Wrapper with fallback when LLM analysis fails."""
    try:
        return analyze_document_for_risks(text, model_name)
    except Exception as e:
        print("⚠️ safe_analyze_document_for_risks fallback triggered:", e)
        risks = {"High": [], "Medium": [], "Low": []}
        obligations = ["Manual review recommended due to analysis limitations"]
        return risks, obligations

# -------------------------------------------------------------------------
# Date Validators & Smart Semantic Search
# -------------------------------------------------------------------------
def simple_date_validator(text: str) -> Dict[str, Any]:
    """Simple date extraction/validity check using ISO-like date patterns."""
    date_matches = re.findall(r"\b\d{4}[-/]\d{2}[-/]\d{2}\b", text)
    results = {"status": "UNKNOWN", "message": "No clear dates found", "action_needed": []}

    if date_matches:
        try:
            expiry = parse_date(date_matches[-1])
            now = datetime.now()
            if expiry < now:
                results["status"] = "EXPIRED"
                results["message"] = f"Expired on {expiry.date()}"
                results["action_needed"].append("Renew or renegotiate the contract")
            elif expiry < now + timedelta(days=30):
                results["status"] = "EXPIRING_SOON"
                results["message"] = f"Expires soon on {expiry.date()}"
                results["action_needed"].append("Plan renewal or extension")
            else:
                results["status"] = "VALID"
                results["message"] = f"Valid until {expiry.date()}"
        except Exception as e:
            results["message"] = f"Date parsing failed: {e}"
    return results

def semantic_search_with_dates(query: str, collection_name: str = "legal_docs", top_k: int = 5) -> str:
    """Run semantic search and appended date-validation result for time-sensitive queries."""
    base_answer = semantic_search(query=query, collection_name=collection_name, top_k=top_k)
    if collection_name not in VECTOR_STORES:
        return base_answer

    vectordb = VECTOR_STORES[collection_name]
    docs = vectordb.similarity_search(query, k=top_k)
    combined_text = " ".join(d.page_content for d in docs)
    validation = simple_date_validator(combined_text)

    if validation["status"] == "EXPIRED":
        return f"🔴 NO - {validation['message']}"
    elif validation["status"] == "EXPIRING_SOON":
        return f"🟡 ALMOST EXPIRED - {validation['message']}"
    elif validation["status"] == "VALID":
        return f"🟢 YES - {validation['message']}"
    else:
        return base_answer

def intelligent_date_extraction_and_validation(text: str) -> Dict[str, Any]:
    """Find multiple ISO-like dates in text and classify them by status."""
    date_matches = re.findall(r"\b\d{4}[-/]\d{2}[-/]\d{2}\b", text)
    extracted_dates = []
    now = datetime.now()
    for d in date_matches:
        try:
            dt = parse_date(d)
            if dt < now:
                status = "Past"
            elif dt < now + timedelta(days=30):
                status = "Expiring Soon"
            else:
                status = "Future"
            extracted_dates.append({"date": str(dt.date()), "status": status})
        except Exception:
            continue
    return {"dates": extracted_dates, "count": len(extracted_dates)}

def semantic_search_with_intelligent_validation(query: str, collection_name: str = "legal_docs", top_k: int = 5) -> str:
    """Semantic search plus a helpful date-extraction summary appended to the LLM answer."""
    base_answer = semantic_search(query=query, collection_name=collection_name, top_k=top_k)
    if collection_name not in VECTOR_STORES:
        return base_answer

    vectordb = VECTOR_STORES[collection_name]
    docs = vectordb.similarity_search(query, k=top_k)
    combined_text = " ".join(d.page_content for d in docs)
    extracted = intelligent_date_extraction_and_validation(combined_text)

    if not extracted["dates"]:
        return base_answer + "\n\nℹ️ No clear dates found in document."
    else:
        summary_lines = [f"- {d['date']} → {d['status']}" for d in extracted["dates"]]
        return base_answer + "\n\n📅 Date Analysis:\n" + "\n".join(summary_lines)

# -------------------------------------------------------------------------
# Utilities
# -------------------------------------------------------------------------
def compare_documents(text1: str, text2: str) -> str:
    """Simple unified-diff style comparison returned as markdown codeblock."""
    a_lines, b_lines = text1.splitlines(), text2.splitlines()
    diff = difflib.unified_diff(a_lines, b_lines, fromfile="Document A", tofile="Document B", lineterm="")
    md_lines = ["### Document Comparison\n", "```diff"]
    md_lines.extend(list(diff))
    md_lines.append("```")
    return "\n".join(md_lines)

def embed_texts(texts: List[str], embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2") -> List[List[float]]:
    """Return raw embeddings for a list of texts (uses HuggingFaceEmbeddings)."""
    embedder = _get_embedding_model(embedding_model_name)
    return embedder.embed_documents(texts)

# -------------------------------------------------------------------------
# Entry point
# -------------------------------------------------------------------------
if __name__ == "__main__":
    print("Backend module loaded (FAISS, OpenRouter).")
    print("Model ID:", MODEL_ID)




